"""
Resume text extraction from PDF and DOCX files.
"""
import io
from pathlib import Path


def extract_text_from_pdf(file_bytes: bytes) -> str:
    """Extract text from PDF bytes using pdfplumber."""
    import pdfplumber

    text_parts = []
    with pdfplumber.open(io.BytesIO(file_bytes)) as pdf:
        for page in pdf.pages:
            page_text = page.extract_text()
            if page_text:
                text_parts.append(page_text)
    return "\n".join(text_parts).strip()


def extract_text_from_docx(file_bytes: bytes) -> str:
    """Extract text from DOCX bytes."""
    from docx import Document

    doc = Document(io.BytesIO(file_bytes))
    paragraphs = []
    for para in doc.paragraphs:
        if para.text.strip():
            paragraphs.append(para.text.strip())
    # Also extract tables
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                if cell.text.strip():
                    paragraphs.append(cell.text.strip())
    return "\n".join(paragraphs)


def extract_text_from_txt(file_bytes: bytes) -> str:
    """Decode plain text."""
    try:
        return file_bytes.decode("utf-8")
    except UnicodeDecodeError:
        return file_bytes.decode("latin-1", errors="ignore")


def extract_resume_text(filename: str, file_bytes: bytes) -> str:
    """Route to correct extractor based on file extension."""
    ext = Path(filename).suffix.lower()
    if ext == ".pdf":
        return extract_text_from_pdf(file_bytes)
    elif ext in (".docx", ".doc"):
        return extract_text_from_docx(file_bytes)
    elif ext == ".txt":
        return extract_text_from_txt(file_bytes)
    else:
        # Try PDF first, then plain text
        try:
            return extract_text_from_pdf(file_bytes)
        except Exception:
            return extract_text_from_txt(file_bytes)
